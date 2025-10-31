# app/services/file_extractor.py
"""
Service for extracting text from various file formats (PDF, DOCX)
"""

import PyPDF2
from docx import Document
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class FileExtractor:
    """Extract text content from uploaded files"""

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text as string
        """
        try:
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from file based on type

        Args:
            file_path: Path to the file
            file_type: MIME type or extension (pdf, docx)

        Returns:
            Extracted text as string
        """
        file_type = file_type.lower()

        if "pdf" in file_type:
            return FileExtractor.extract_from_pdf(file_path)
        elif "docx" in file_type or "document" in file_type:
            return FileExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
